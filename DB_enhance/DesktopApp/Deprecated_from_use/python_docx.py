import docx   # Access the Python docx module
from docx.shared import Pt   # Allows for changing font size
from docx.shared import Cm   # Allows for changing table row height in centimeters
from docx.oxml import OxmlElement   # Needed for adjusting table cell margins
from docx.oxml.ns import qn  # Needed for adjusting table cell margins
def python_docx():
    print("test")
    # Shane Snediker, Ethan Wolcott, Trevor Troxel, Pragalva Dhungana and Pukar Mahat
    # CS 472 Software Engineering
    # Hewitt Research Foundation database rehaul project
    # File last updated: 5-8-2021

    # This file contains the work that we did in pursuing the Python Docx library to
    # attempt to put together a print template that we could use to print student 
    # information onto scantron tests.  Unfortunately, the Word Doc formatting did not
    # provide enough flexibility in spacing out the elements that Hewitt needs to print
    # onto student tests, so we had to pursue an alternative solution

    # But we've left this file here for possible future reference.  The Python Docx library
    # allows you to use the Python language to generate a Word Document



    ################## FUNCTION DEFINITIONS ################################################

    # A function that iterates through the cells of a Word doc table
    # and applies bold font to the text within the cells
    # param: the docx table column whose text needs to be boldened
    # return: void
    def make_column_bold(*columns):
        # columns is an array of table column cells, here we step
        # through them 1 by 1
        for column in columns:
            # And focus in on each cell
            for cell in column.cells:
                # And specifically the paragraphs in the cell
                for paragraph in cell.paragraphs:
                    # For Python docx, all text with altered formatting
                    # gets assigned a separate "run".  So, in order to 
                    # modify the table column font, we access it through
                    # its run
                    for run in paragraph.runs:
                        run.font.bold = True

    # This function takes in a specific table cell and a set of specified margin parameters and 
    # alters the width of the cell margins to fit the passed in parameters
    # param: cell : the cell instance that needs to have its margins modified
    #        **kwargs : a choice of which sides of the cell to modify (top, bottom, start, end)
    #                   and the specified margin value (provided values are 1/1440 of an inch)
    # return: void
    # Example: set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
    # Read more here: http://officeopenxml.com/WPtableCellMargins.php
    def set_cell_margins(cell, **kwargs):
        
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        for m in [
            "top",
            "start",
            "bottom",
            "end",
        ]:
            if m in kwargs:
                node = OxmlElement("w:{}".format(m))
                node.set(qn('w:w'), str(kwargs.get(m)))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)

        tcPr.append(tcMar)

    #############  CREATE AND DEFINE WORD DOC  ################################################

    # Instantiate the Word document
    scantron_template = docx.Document()

    # Set font style and size
    Tstyle = scantron_template.styles['Normal']
    font = Tstyle.font
    font.name = 'Calibri'
    font.size = Pt(10)

    ########################  STUDENT INFORMATION SECTION ####################################

    # This section defines the student information section that will be printed in the upper
    # left area of the scantron sheets.  Here we will use a table to organize the student 
    # information.  Our table will be 1 row by 2 columns.  The left column will hold the 
    # data headers while the right column will hold the specific corresponding student info.
    # Within the row of information will be attributes separated by an end of line (in this order): 
    # Group ID (if applicable), Student ID, Student name, Student grade,Test print date, Parent 
    # information 

    # Create the 1X2 table to hold the student information
    student_info = scantron_template.add_table(rows = 1, cols = 2)

    # Python docx gives definition to each cell within a document table.
    # The student table cells are defined in the following way:
    #                 COLUMN 0    COLUMN 1            
    #   ROW 0:     |   (0,0)   |   (0,1)   |

    #              | Group Id     | Student's group id if part of a group |
    #              | Student Id   | Student's id number                   |
    #              | Student      | Student's first and last name         |
    #              | Grade        | Grade level of student                |
    #              | Date Printed | Date the test was printed             |
    #              | Parents      | Student's parents' name and address   |

    # Let's adjust the margins on the left side of the document and the top of the document
    # so that we can place the student info right where we want it in the upper left corner

    margin_adjust = scantron_template.sections  # Access document sections
    for section in margin_adjust:               # Iterate through each section of the document 
        section.left_margin = Cm(1.65)          # Set the left margin to 1.65 centimeters
        section.top_margin = Cm(2)              # Set the top marging to 2 centimeters

    # Let's specify the width of the table cells so Hewitt's student
    # information will print up nicely
    # NOTE: Width of cell objects is measured in EMU's.  914,400 EMU's = 1 inch
    # Here we make column 1 (the attribute headers) roughly 1.2 inches wide and
    # column 2 (the actual student data) roughly 2.5 inches wide
    student_info.cell(0,0).width = 914400     # 1 * 914400 = 914,400
    student_info.cell(0,1).width = 2286000     # 2.5 * 914400 =2,286,000

    # Now let's make the student info table row high enough to contain
    # all of the pieces of pertinent student information (5 centimeters)
    student_info.rows[0].height = Cm(5)

    # Now we are ready to populate the table with data

    # Let's give the student information table its entries
    # First the left column which is the attribute headers
    student_info.cell(0, 0).text = 'Group ID:\nStudent ID:\nStudent:\nGrade:\nDate Printed:\nParents:'

    # Don't forget to make the attribute headers bold by calling the 
    # make_column_bold function and passing it column 0
    make_column_bold(student_info.columns[0])

    # Now we can populate the attribute fields with this student's information

    # ENTER STUDENT DATA HERE
    student_info.cell(0, 1).text = '12345 \n 1 \n James Richardson \n 7 \n 02-07-2020 \n James and Sally Richardson \n N. 15 Hewitt Ln. \n Madison WI, 95678'       

    #######################  NON TRANSFERABLE PARAGRAPH SECTION #################################

    # Add not transferable statement paragraph to the document and then make it bold
    not_transferable = scantron_template.add_paragraph()
    not_transferable.add_run("Not Transferrable To Any Other Parent or Student").bold = True

    ####################### GRADE DOTS SECTION ##################################################

    # Let's instantiate a table that will hold the configuration of grade dots 
    grade_dot = scantron_template.add_table(rows = 1, cols = 2)
    # Using 2 columns for this table allows us to horizontally align the grade column
    # where we need it to be so that the grade dots will land in the correct circle
    # From a series of guess and checks, we deduced that the grade dot column (grade_dot.cell(0,1))
    # needs to start 4.4 centimeters from the left margin of the Word doc
    grade_dot.cell(0,0).width = Cm(4.4)
    grade_dot.cell(0,1).width = Cm(1.0)
    # Set the column that grade dots will be placed in to an appropriate height
    grade_dot.rows[0].height = Cm(4.0)

    # In order to get the scantron grade bubble filled in, we've generated 6
    # png images of a black dot spaced out in different configurations.  To 
    # get accurate placements of the grade dots based on the specific grade 
    # of the student, we load a different image for each grade.

    # First we load a paragraph element into the grade_dot column
    grade_dot_p = grade_dot.cell(0,1).add_paragraph()
    # Pictures must be loaded into a run
    grade_dot_r = grade_dot_p.add_run()
    # Add the appropriate grade image into the run
    grade_dot_r.add_picture('img/grade3_circle.PNG', width = Cm(2), height = Cm(4))

    ####################### FOR OFFICE USE ONLY SECTION #########################################

    # Initialize the table that will hold the office use only rectangular markings
    office_use = scantron_template.add_table(rows = 7, cols = 1)

    # Remove the margins for this table because the lines in the scantron "for office use only"
    # section are really close together and we're going to need our rectangles to not be vertically
    # separated by much space
    set_cell_margins(office_use.cell(0,0), top=25, bottom=0)  # Remove row 0's bottom margin
    set_cell_margins(office_use.cell(1,0), top=0, bottom=0)  # Remove row 1's bottom and top margin
    set_cell_margins(office_use.cell(2,0), top=0, bottom=0)  # Remove row 2's bottom and top margin
    set_cell_margins(office_use.cell(3,0), top=0, bottom=0)  # Remove row 3's bottom and top margin
    set_cell_margins(office_use.cell(4,0), top=0, bottom=0)  # Remove row 4's bottom and top margin
    set_cell_margins(office_use.cell(5,0), top=0, bottom=0)  # Remove row 5's bottom and top margin
    set_cell_margins(office_use.cell(6,0), top=0, bottom=0)  # Remove row 6's bottom and top margin

    # Now we specify the width of the rows -> 4 inches * 914400 = 3,657,600
    # And the height of each row to 1/2 centimeter
    for row in range(0, len(student_info.rows)):
        office_use.rows[row].width = 3657600
        office_use.rows[row].height = Cm(0.5)

    # Now let's populate the "for office use only" section with the appropriate rectangles

    account_box_p = office_use.cell(0,0).add_paragraph()
    account_box_r = account_box_p.add_run()
    account_box_r.add_picture('img/rectangle.PNG', width = Cm(.5), height = Cm(.1))

    account_box_p = office_use.cell(1,0).add_paragraph()
    account_box_r = account_box_p.add_run()
    account_box_r.add_picture('img/rectangle.PNG', width = Cm(.5), height = Cm(.1))

    account_box_p = office_use.cell(2,0).add_paragraph()
    account_box_r = account_box_p.add_run()
    account_box_r.add_picture('img/rectangle.PNG', width = Cm(.5), height = Cm(.1))

    account_box_p = office_use.cell(3,0).add_paragraph()
    account_box_r = account_box_p.add_run()
    account_box_r.add_picture('img/rectangle.PNG', width = Cm(.5), height = Cm(.1))

    # Now we can merge the rows together to get the rectangles even closer together

    # NOTE: Remember the zero indexing factor.  For example, when I say "Merge the 1st
    #       and 2nd row together" below, the first row in the system is defined as row
    #       zero (or cell (0,0)), whereas the 2nd row in the system is defined as row
    #       one (or cell (1,0)).  Just a reminder.

    # Merge the 1st and 2nd row together
    row_one_bottom = office_use.cell(0,0)
    row_two_top = office_use.cell(1,0)
    row0_row1 = row_one_bottom.merge(row_two_top)

    # Merge the 2nd and 3rd row together
    row_two_bottom = office_use.cell(1,0)
    row_three_top = office_use.cell(2,0)
    row1_row2 = row_two_bottom.merge(row_three_top)

    # Merge the 3rd and 4th row together
    row_three_bottom = office_use.cell(2,0)
    row_four_top = office_use.cell(3,0)
    row2_row3 = row_three_bottom.merge(row_four_top)

    # Merge the 4th and 5th row together
    row_four_bottom = office_use.cell(3,0)
    row_five_top = office_use.cell(4,0)
    row3_row4 = row_four_bottom.merge(row_five_top)

    # Merge the 5th and 6th row together
    row_five_bottom = office_use.cell(4,0)
    row_six_top = office_use.cell(5,0)
    row4_row5 = row_five_bottom.merge(row_six_top)

    # Merge the 6th and 7th row together
    row_six_bottom = office_use.cell(5,0)
    row_seven_top = office_use.cell(6,0)
    row5_row6 = row_six_bottom.merge(row_seven_top)

    #############################################################################################

    # Save the document
    scantron_template.save('test.docx')

